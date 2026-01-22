"""
Text Extraction Pipeline for Doctor's Notes Attachments.

Supports:
- PDF: PyMuPDF (fitz) for text extraction
- DOCX: python-docx for structured text
- Images: Tesseract OCR (pytesseract)
- TXT: Direct reading with encoding detection

All extraction is async-safe and includes error handling.
"""

import asyncio
import hashlib
import io
import logging
import re
from abc import ABC, abstractmethod
from pathlib import Path
from typing import NamedTuple, Optional

logger = logging.getLogger(__name__)


class ExtractionResult(NamedTuple):
    """Result of text extraction."""

    text: str
    page_count: Optional[int] = None
    word_count: int = 0
    char_count: int = 0
    image_width: Optional[int] = None
    image_height: Optional[int] = None
    error: Optional[str] = None

    @property
    def success(self) -> bool:
        return self.error is None


class BaseExtractor(ABC):
    """Base class for text extractors."""

    @abstractmethod
    async def extract(self, content: bytes, filename: str) -> ExtractionResult:
        """Extract text from file content."""
        pass

    @staticmethod
    def count_words(text: str) -> int:
        """Count words in text."""
        return len(text.split())

    @staticmethod
    def clean_text(text: str) -> str:
        """Clean extracted text."""
        # Normalize whitespace
        text = re.sub(r"\s+", " ", text)
        # Remove control characters except newlines
        text = re.sub(r"[\x00-\x08\x0b\x0c\x0e-\x1f\x7f-\x9f]", "", text)
        return text.strip()


class PDFExtractor(BaseExtractor):
    """Extract text from PDF files using PyMuPDF."""

    async def extract(self, content: bytes, filename: str) -> ExtractionResult:
        try:
            import fitz  # PyMuPDF

            # Run in thread pool to avoid blocking
            loop = asyncio.get_event_loop()
            result = await loop.run_in_executor(None, self._extract_sync, content)
            return result
        except ImportError:
            return ExtractionResult(
                text="",
                error="PyMuPDF not installed. Install with: pip install pymupdf",
            )
        except Exception as e:
            logger.exception(f"PDF extraction failed for {filename}")
            return ExtractionResult(text="", error=str(e))

    def _extract_sync(self, content: bytes) -> ExtractionResult:
        import fitz

        doc = fitz.open(stream=content, filetype="pdf")
        text_parts = []
        page_count = len(doc)

        for page_num in range(page_count):
            page = doc[page_num]
            text = page.get_text("text")
            if text.strip():
                text_parts.append(f"[Page {page_num + 1}]\n{text}")

        doc.close()

        full_text = self.clean_text("\n\n".join(text_parts))
        return ExtractionResult(
            text=full_text,
            page_count=page_count,
            word_count=self.count_words(full_text),
            char_count=len(full_text),
        )


class DOCXExtractor(BaseExtractor):
    """Extract text from DOCX files using python-docx."""

    async def extract(self, content: bytes, filename: str) -> ExtractionResult:
        try:
            from docx import Document

            loop = asyncio.get_event_loop()
            result = await loop.run_in_executor(None, self._extract_sync, content)
            return result
        except ImportError:
            return ExtractionResult(
                text="",
                error="python-docx not installed. Install with: pip install python-docx",
            )
        except Exception as e:
            logger.exception(f"DOCX extraction failed for {filename}")
            return ExtractionResult(text="", error=str(e))

    def _extract_sync(self, content: bytes) -> ExtractionResult:
        from docx import Document

        doc = Document(io.BytesIO(content))
        text_parts = []

        # Extract paragraphs
        for para in doc.paragraphs:
            if para.text.strip():
                text_parts.append(para.text)

        # Extract tables
        for table in doc.tables:
            table_text = []
            for row in table.rows:
                row_text = [cell.text.strip() for cell in row.cells]
                table_text.append(" | ".join(row_text))
            if table_text:
                text_parts.append("\n".join(table_text))

        full_text = self.clean_text("\n\n".join(text_parts))
        return ExtractionResult(
            text=full_text,
            word_count=self.count_words(full_text),
            char_count=len(full_text),
        )


class ImageExtractor(BaseExtractor):
    """Extract text from images using Tesseract OCR."""

    # Languages for OCR (English and Greek for Cyprus)
    LANGUAGES = "eng+ell"

    async def extract(self, content: bytes, filename: str) -> ExtractionResult:
        try:
            from PIL import Image
            import pytesseract

            loop = asyncio.get_event_loop()
            result = await loop.run_in_executor(None, self._extract_sync, content)
            return result
        except ImportError as e:
            missing = "PIL" if "PIL" in str(e) else "pytesseract"
            return ExtractionResult(
                text="",
                error=f"{missing} not installed. Install with: pip install pillow pytesseract",
            )
        except Exception as e:
            logger.exception(f"Image OCR failed for {filename}")
            return ExtractionResult(text="", error=str(e))

    def _extract_sync(self, content: bytes) -> ExtractionResult:
        from PIL import Image
        import pytesseract

        image = Image.open(io.BytesIO(content))
        width, height = image.size

        # Perform OCR
        text = pytesseract.image_to_string(image, lang=self.LANGUAGES)
        text = self.clean_text(text)

        return ExtractionResult(
            text=text,
            word_count=self.count_words(text),
            char_count=len(text),
            image_width=width,
            image_height=height,
        )


class TextExtractor(BaseExtractor):
    """Extract text from plain text files."""

    ENCODINGS = ["utf-8", "utf-16", "latin-1", "cp1253"]  # cp1253 for Greek

    async def extract(self, content: bytes, filename: str) -> ExtractionResult:
        text = None

        # Try different encodings
        for encoding in self.ENCODINGS:
            try:
                text = content.decode(encoding)
                break
            except UnicodeDecodeError:
                continue

        if text is None:
            return ExtractionResult(
                text="",
                error="Unable to decode text file with known encodings",
            )

        text = self.clean_text(text)
        return ExtractionResult(
            text=text,
            word_count=self.count_words(text),
            char_count=len(text),
        )


class ExtractionService:
    """
    Main service for extracting text from attachments.

    Automatically selects the appropriate extractor based on file type.
    """

    # Supported MIME types and their extractors
    EXTRACTORS: dict[str, type[BaseExtractor]] = {
        # PDF
        "application/pdf": PDFExtractor,
        # Word documents
        "application/vnd.openxmlformats-officedocument.wordprocessingml.document": DOCXExtractor,
        "application/msword": DOCXExtractor,  # .doc (basic support via docx)
        # Images
        "image/jpeg": ImageExtractor,
        "image/png": ImageExtractor,
        "image/tiff": ImageExtractor,
        "image/bmp": ImageExtractor,
        "image/webp": ImageExtractor,
        # Text
        "text/plain": TextExtractor,
        "text/csv": TextExtractor,
        "text/html": TextExtractor,
    }

    # File type classification
    FILE_TYPES = {
        "application/pdf": "pdf",
        "application/vnd.openxmlformats-officedocument.wordprocessingml.document": "docx",
        "application/msword": "doc",
        "image/jpeg": "image",
        "image/png": "image",
        "image/tiff": "image",
        "image/bmp": "image",
        "image/webp": "image",
        "text/plain": "txt",
        "text/csv": "txt",
        "text/html": "txt",
    }

    # Max file sizes (bytes)
    MAX_SIZES = {
        "pdf": 50 * 1024 * 1024,  # 50MB
        "docx": 50 * 1024 * 1024,
        "doc": 50 * 1024 * 1024,
        "image": 20 * 1024 * 1024,  # 20MB for images
        "txt": 10 * 1024 * 1024,  # 10MB for text
    }

    def __init__(self) -> None:
        self._extractors: dict[str, BaseExtractor] = {}

    def get_extractor(self, mime_type: str) -> Optional[BaseExtractor]:
        """Get or create extractor for MIME type."""
        if mime_type not in self.EXTRACTORS:
            return None

        if mime_type not in self._extractors:
            self._extractors[mime_type] = self.EXTRACTORS[mime_type]()

        return self._extractors[mime_type]

    def get_file_type(self, mime_type: str) -> Optional[str]:
        """Get file type classification from MIME type."""
        return self.FILE_TYPES.get(mime_type)

    def is_supported(self, mime_type: str) -> bool:
        """Check if MIME type is supported."""
        return mime_type in self.EXTRACTORS

    def validate_size(self, content: bytes, mime_type: str) -> Optional[str]:
        """Validate file size. Returns error message if invalid."""
        file_type = self.get_file_type(mime_type)
        if not file_type:
            return f"Unsupported file type: {mime_type}"

        max_size = self.MAX_SIZES.get(file_type, 10 * 1024 * 1024)
        if len(content) > max_size:
            return f"File too large. Maximum size for {file_type}: {max_size // (1024 * 1024)}MB"

        return None

    async def extract(
        self, content: bytes, filename: str, mime_type: str
    ) -> ExtractionResult:
        """
        Extract text from file content.

        Args:
            content: Raw file bytes
            filename: Original filename
            mime_type: MIME type of the file

        Returns:
            ExtractionResult with text and metadata
        """
        # Validate size
        size_error = self.validate_size(content, mime_type)
        if size_error:
            return ExtractionResult(text="", error=size_error)

        # Get extractor
        extractor = self.get_extractor(mime_type)
        if not extractor:
            return ExtractionResult(
                text="",
                error=f"No extractor available for {mime_type}",
            )

        # Perform extraction
        return await extractor.extract(content, filename)

    @staticmethod
    def compute_checksum(content: bytes) -> str:
        """Compute SHA-256 checksum of content."""
        return hashlib.sha256(content).hexdigest()

    @staticmethod
    def detect_mime_type(content: bytes, filename: str) -> Optional[str]:
        """
        Detect MIME type from content and filename.

        Uses python-magic for content-based detection with filename fallback.
        """
        try:
            import magic

            # Content-based detection
            mime = magic.from_buffer(content, mime=True)
            if mime and mime != "application/octet-stream":
                return mime
        except ImportError:
            logger.warning("python-magic not installed, using filename-based detection")
        except Exception:
            pass

        # Filename-based fallback
        ext = Path(filename).suffix.lower()
        extension_map = {
            ".pdf": "application/pdf",
            ".docx": "application/vnd.openxmlformats-officedocument.wordprocessingml.document",
            ".doc": "application/msword",
            ".jpg": "image/jpeg",
            ".jpeg": "image/jpeg",
            ".png": "image/png",
            ".tiff": "image/tiff",
            ".tif": "image/tiff",
            ".bmp": "image/bmp",
            ".webp": "image/webp",
            ".txt": "text/plain",
            ".csv": "text/csv",
            ".html": "text/html",
            ".htm": "text/html",
        }
        return extension_map.get(ext)


# Singleton instance
extraction_service = ExtractionService()
